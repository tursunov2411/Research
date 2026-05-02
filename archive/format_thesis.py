import sys
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def main():
    doc = docx.Document("raw_report.docx")
    
    # 1. PAGE SETUP
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7) # A4
        
    # 2. DEFAULT TYPOGRAPHY
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5

    # Prepare new document to hold the restructured content
    new_doc = docx.Document()
    for section in new_doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        
    normal_style = new_doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.line_spacing = 1.5

    # Helper function to add major heading
    def add_major_heading(text, page_break_before=True):
        if page_break_before and len(new_doc.paragraphs) > 0:
            new_doc.add_page_break()
        p = new_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # 3. DOCUMENT STRUCTURE placeholders
    
    # 1. Title Page
    p = new_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI\n"
                    "TOSHKENT DAVLAT IQTISODIYOT UNIVERSITETI\n\n"
                    "[TITLE PAGE - CONTENT REQUIRED]")
    run.font.size = Pt(14)
    
    # 2. Back of title page (signatures)
    new_doc.add_page_break()
    new_doc.add_paragraph("[BACK OF TITLE PAGE — SIGNATURES BLOCK REQUIRED]")

    # 3. Annotation in Uzbek
    new_doc.add_page_break()
    add_major_heading("ANNOTATSIYA", False)
    new_doc.add_paragraph("[SECTION MISSING — CONTENT REQUIRED]")
    p = new_doc.add_paragraph()
    p.add_run("Kalit so'zlar: ").bold = True
    p.add_run("[KEYWORDS IN UZBEK REQUIRED]").italic = True

    # 4. Annotation in English
    new_doc.add_page_break()
    add_major_heading("ANNOTATION", False)
    # Check if there's an abstract in the original document
    abstract_text = ""
    for p_old in doc.paragraphs:
        if p_old.text.lower().startswith("this paper examines uzbekistan's"):
            abstract_text = p_old.text
            break
    
    if abstract_text:
        new_doc.add_paragraph(abstract_text)
    else:
        new_doc.add_paragraph("[SECTION MISSING — CONTENT REQUIRED]")
        
    p = new_doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("Global value chains, FDI-led industrialisation, Uzbekistan, Northeast Asia, revealed comparative advantage, economic complexity, structural break").italic = True

    # 6. Table of Contents
    add_major_heading("MUNDARIJA")
    new_doc.add_paragraph("[TABLE OF CONTENTS AUTO-GENERATED OR INSERTED HERE]")

    # 7. List of Tables
    add_major_heading("KELTIRILGAN JADVALLAR RO'YXATI")
    new_doc.add_paragraph("[LIST OF TABLES HERE]")

    # 8. List of Figures
    add_major_heading("KELTIRILGAN RASMLAR RO'YXATI")
    new_doc.add_paragraph("[LIST OF FIGURES HERE]")

    # 9. List of Abbreviations (if applicable)
    add_major_heading("KELTIRILGAN QISQARTMALAR")
    new_doc.add_paragraph("[LIST OF ABBREVIATIONS HERE]")

    # We will iterate through the original document and place content into chapters
    current_chapter = None
    
    # Let's map sections from original doc to chapters
    # 10. CHAPTER I — INTRODUCTION (Kirish)
    add_major_heading("I BOB. KIRISH (INTRODUCTION)")
    
    intro_found = False
    lit_review_found = False
    method_found = False
    analysis_found = False
    conclusion_found = False
    
    chapter_num = 1
    
    for p_old in doc.paragraphs:
        text = p_old.text.strip()
        if not text:
            continue
            
        lower_text = text.lower()
        
        # Check for headings to map them to correct Chapters
        if "introduction" in lower_text and len(text.split()) < 5:
            current_chapter = "I BOB"
            intro_found = True
            continue
        elif "literature review" in lower_text and len(text.split()) < 10:
            add_major_heading("II BOB. ADABIYOTLAR SHARHI (LITERATURE REVIEW)")
            current_chapter = "II BOB"
            lit_review_found = True
            continue
        elif "methodology" in lower_text and len(text.split()) < 10:
            add_major_heading("III BOB. METODOLOGIYA (METHODOLOGY)")
            current_chapter = "III BOB"
            method_found = True
            continue
        elif ("analysis" in lower_text or "results" in lower_text or "structural break" in lower_text or "regression" in lower_text) and len(text.split()) < 15 and current_chapter != "IV BOB" and not "literature" in lower_text:
            if not analysis_found:
                add_major_heading("IV BOB. TAHLIL VA NATIJALAR (ANALYSIS AND RESULTS)")
                current_chapter = "IV BOB"
                analysis_found = True
            # Keep the old heading as a subsection
            p_new = new_doc.add_paragraph()
            p_new.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p_new.add_run(text)
            r.bold = True
            continue
        elif "conclusion" in lower_text and len(text.split()) < 5:
            add_major_heading("V BOB. XULOSA (CONCLUSION)")
            current_chapter = "V BOB"
            conclusion_found = True
            continue
        elif "references" in lower_text or "bibliography" in lower_text:
            current_chapter = "REFERENCES"
            add_major_heading("FOYDALANILGAN ADABIYOTLAR RO'YXATI")
            continue
        elif "appendices" in lower_text or "appendix" in lower_text:
            current_chapter = "APPENDICES"
            add_major_heading("ILOVALAR (APPENDICES)")
            continue
            
        if current_chapter:
            # Figure/Table conversions
            # Figure 1: -> 1-rasm.
            if lower_text.startswith("figure"):
                m = re.match(r'figure\s+(\d+)[\.:\s]+(.*)', text, re.IGNORECASE)
                if m:
                    num = m.group(1)
                    caption = m.group(2)
                    
                    p_new = new_doc.add_paragraph()
                    p_new.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p_new.add_run(f"{num}-rasm. {caption}")
                    r.bold = True
                    
                    if "source" not in caption.lower() and "manba" not in caption.lower():
                        p_new = new_doc.add_paragraph()
                        r = p_new.add_run("[CAPTION/SOURCE REQUIRED]")
                        r.font.size = Pt(10)
                        r.italic = True
                    continue
            
            # Table 1: -> 1-jadval
            if lower_text.startswith("table"):
                m = re.match(r'table\s+(\d+)[\.:\s]+(.*)', text, re.IGNORECASE)
                if m:
                    num = m.group(1)
                    title = m.group(2)
                    
                    p_num = new_doc.add_paragraph()
                    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p_num.add_run(f"{num}-jadval")
                    
                    p_title = new_doc.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p_title.add_run(title)
                    r.bold = True
                    
                    if "source" not in title.lower() and "manba" not in title.lower():
                        p_new = new_doc.add_paragraph()
                        r = p_new.add_run("[CAPTION/SOURCE REQUIRED]")
                        r.font.size = Pt(10)
                        r.italic = True
                    continue

            # Standard paragraph
            p_new = new_doc.add_paragraph(text)
            # if paragraph seems like a section heading, format it
            if p_old.style.name.startswith('Heading'):
                p_new.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p_new.runs:
                    r.bold = True

    if not intro_found:
        p = new_doc.add_paragraph("[SECTION MISSING — CONTENT REQUIRED]")
    if not lit_review_found:
        add_major_heading("II BOB. ADABIYOTLAR SHARHI (LITERATURE REVIEW)")
        p = new_doc.add_paragraph("[SECTION MISSING — CONTENT REQUIRED]")
    if not method_found:
        add_major_heading("III BOB. METODOLOGIYA (METHODOLOGY)")
        p = new_doc.add_paragraph("[SECTION MISSING — CONTENT REQUIRED]")
    if not analysis_found:
        add_major_heading("IV BOB. TAHLIL VA NATIJALAR (ANALYSIS AND RESULTS)")
        p = new_doc.add_paragraph("[SECTION MISSING — CONTENT REQUIRED]")
    if not conclusion_found:
        add_major_heading("V BOB. XULOSA (CONCLUSION)")
        p = new_doc.add_paragraph("[SECTION MISSING — CONTENT REQUIRED]")

    new_doc.save('formatted_thesis.docx')
    print("Formatting complete. Output saved to formatted_thesis.docx")

if __name__ == "__main__":
    main()
